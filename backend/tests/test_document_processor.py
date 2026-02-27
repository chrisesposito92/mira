"""Tests for document text extraction and processing pipeline."""

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import uuid4

import pytest

from app.services.document_processor import (
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_txt,
    process_document,
)


class TestExtractTextFromTxt:
    def test_reads_text_file(self, tmp_path: Path):
        f = tmp_path / "test.txt"
        f.write_text("Hello, world!\nLine two.", encoding="utf-8")
        result = extract_text_from_txt(f)
        assert "Hello, world!" in result
        assert "Line two." in result


class TestExtractTextFromPdf:
    def test_reads_pdf(self, tmp_path: Path):
        """Create a minimal PDF and extract text."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        pdf_path = tmp_path / "test.pdf"
        with open(pdf_path, "wb") as f:
            writer.write(f)
        result = extract_text_from_pdf(pdf_path)
        assert isinstance(result, str)


class TestExtractTextFromDocx:
    def test_reads_docx(self, tmp_path: Path):
        """Create a minimal DOCX and extract text."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test paragraph one.")
        doc.add_paragraph("Test paragraph two.")
        docx_path = tmp_path / "test.docx"
        doc.save(docx_path)

        result = extract_text_from_docx(docx_path)
        assert "Test paragraph one." in result
        assert "Test paragraph two." in result


def _make_mock_supabase():
    """Create a mock Supabase client with chainable table methods."""
    mock = MagicMock()
    mock_builder = MagicMock()
    mock_builder.update = MagicMock(return_value=mock_builder)
    mock_builder.eq = MagicMock(return_value=mock_builder)
    mock_builder.execute = MagicMock()
    mock.table = MagicMock(return_value=mock_builder)
    return mock


class TestProcessDocument:
    async def test_successful_processing(self, tmp_path: Path):
        """Test full pipeline: extract -> chunk -> embed -> store."""
        f = tmp_path / "test.txt"
        f.write_text("This is test content for embedding.", encoding="utf-8")

        mock_pool = AsyncMock()
        mock_supabase = _make_mock_supabase()
        doc_id = uuid4()
        project_id = uuid4()

        with patch(
            "app.services.document_processor.ingest_document",
            AsyncMock(return_value=3),
        ):
            count = await process_document(
                pool=mock_pool,
                supabase=mock_supabase,
                document_id=doc_id,
                file_path=f,
                file_type="txt",
                project_id=project_id,
            )

        assert count == 3
        # Verify status was updated to "processing" then "ready"
        update_calls = mock_supabase.table.return_value.update.call_args_list
        assert len(update_calls) >= 2
        assert update_calls[0].args[0] == {"processing_status": "processing"}
        assert update_calls[1].args[0] == {"processing_status": "ready", "chunk_count": 3}

    async def test_failed_processing_sets_error(self, tmp_path: Path):
        """Test that failures update status to 'failed' with error message."""
        f = tmp_path / "test.txt"
        f.write_text("Content", encoding="utf-8")

        mock_pool = AsyncMock()
        mock_supabase = _make_mock_supabase()
        doc_id = uuid4()

        with patch(
            "app.services.document_processor.ingest_document",
            AsyncMock(side_effect=RuntimeError("embedding API error")),
        ):
            with pytest.raises(RuntimeError, match="embedding API error"):
                await process_document(
                    pool=mock_pool,
                    supabase=mock_supabase,
                    document_id=doc_id,
                    file_path=f,
                    file_type="txt",
                    project_id=uuid4(),
                )

        # Verify status was set to failed
        update_calls = mock_supabase.table.return_value.update.call_args_list
        last_update = update_calls[-1].args[0]
        assert last_update["processing_status"] == "failed"
        assert "embedding API error" in last_update["error_message"]

    async def test_unsupported_file_type(self, tmp_path: Path):
        f = tmp_path / "test.xyz"
        f.write_text("Content", encoding="utf-8")

        mock_pool = AsyncMock()
        mock_supabase = _make_mock_supabase()

        with pytest.raises(ValueError, match="Unsupported file type"):
            await process_document(
                pool=mock_pool,
                supabase=mock_supabase,
                document_id=uuid4(),
                file_path=f,
                file_type="xyz",
                project_id=uuid4(),
            )

    async def test_empty_content_raises(self, tmp_path: Path):
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")

        mock_pool = AsyncMock()
        mock_supabase = _make_mock_supabase()

        with pytest.raises(ValueError, match="No text content"):
            await process_document(
                pool=mock_pool,
                supabase=mock_supabase,
                document_id=uuid4(),
                file_path=f,
                file_type="txt",
                project_id=uuid4(),
            )
